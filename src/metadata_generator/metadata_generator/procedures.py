from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

import re
import unicodedata
from docx import Document
import pandas as pd
from typing import Tuple
import numpy as np

from aind_data_schema.core.procedures import Procedures
from aind_data_schema.components.subject_procedures import Surgery
from aind_data_schema_models.units import MassUnit, TimeUnit
from aind_data_schema.components.surgery_procedures import (
    Anaesthetic,
    Craniotomy,
    Headframe,
    ProbeImplant,
    CraniotomyType,
    ProtectiveMaterial,
)
from aind_data_schema.components.coordinates import CoordinateSystemLibrary


# =============================================================================
# ----------------------- Label maps & regexes ---------------------------------
# =============================================================================

# Exact-label map for metadata table keys
KNOWN_FIELDS = {
    "Date / Experiment": "date_experiment",
    "Goal": "goal",
    "Overall impressions / conclusions": "overall_impressions",
    "Animal ID": "animal_id",
    "Animal state / anesthesia induction time": "anesthesia_induction",
    "Heating pad (state/temp)": "heating_pad",
    "Animal birthdate": "animal_birthdate",
    "Animal weight before": "weight_before",
    "Animal weight after": "weight_after",
    "Animal surgery history": "surgery_history",
    "Craniotomy coordinates": "craniotomy_coordinates",
    "Targettng Modifications": "targeting_modifications",  # keep label as-is in doc
    "Final probe depth (from pia)": "final_probe_depth_um",
    "Z(GNDscrew-REFscrew)": "z_gnd_screw_ref_screw",
    "Z(GNDprobe-REFprobe)": "z_gnd_probe_ref_probe",
    "Probe SN": "probe_sn",
    "Probe GND/REF leads connected to": "probe_gnd_ref_connected_to",
    "Probe orientation vs. AP axis": "probe_orientation_ap",
    "Number of probe channels": "probe_channels",
}

# Canonical map for resilient header matching
KNOWN_FIELDS_CANON = {
    "date experiment": "date_experiment",
    "goal": "goal",
    "overall impressions conclusions": "overall_impressions",
    "animal id": "animal_id",
    "animal state anesthesia induction time": "anesthesia_induction",
    "heating pad state temp": "heating_pad",
    "animal birthdate": "animal_birthdate",
    "animal weight before": "weight_before",
    "animal weight after": "weight_after",
    "animal surgery history": "surgery_history",
    "craniotomy coordinates": "craniotomy_coordinates",
    "targettng modifications": "targeting_modifications",
    "final probe depth from pia": "final_probe_depth_um",
    "z gndscrew refscrew": "z_gnd_screw_ref_screw",
    "z gndprobe refprobe": "z_gnd_probe_ref_probe",
    "probe sn": "probe_sn",
    "probe gnd ref leads connected to": "probe_gnd_ref_connected_to",
    "probe orientation vs ap axis": "probe_orientation_ap",
    "number of probe channels": "probe_channels",
}

# Notes headers (paragraphs or table rows, canonicalized)
NOTE_HEADERS_CANON = {
    "surgery notes": "surgery_notes",
    "metabond": "metabond_notes",
    "probe notes": "probe_notes",
    "penetration coverage notes": "penetration_notes",  # "penetration/coverage notes" -> canon removes '/'
    "implantation notes": "implantation_notes",
    "plug in 8 days after implant": "plugin_notes",  # "Plug-in" -> canon removes punctuation
}

# Regex patterns to capture inline note content (group 1 = trailing content)
NOTE_HEADER_PATTERNS = {
    "surgery_notes": re.compile(r"^\s*Surgery\s*notes\s*:?\s*(.*)$", re.IGNORECASE),
    "metabond_notes": re.compile(r"^\s*Metabond\s*:?\s*(.*)$", re.IGNORECASE),
    "probe_notes": re.compile(r"^\s*Probe\s*Notes\s*:?\s*(.*)$", re.IGNORECASE),
    "penetration_notes": re.compile(
        r"^\s*Penetration\s*(?:/|and\s+)?\s*coverage\s*notes\s*:?\s*(.*)$",
        re.IGNORECASE,
    ),
    "implantation_notes": re.compile(r"^\s*Implantation\s*notes\s*:?\s*(.*)$", re.IGNORECASE),
    "plugin_notes": re.compile(
        r"^\s*Plug[-\s]*in\s*8\s*days\s*after\s*implant\s*:?\s*(.*)$", re.IGNORECASE
    ),
}

# Date/experiment lines accept '/', en dash, or hyphen separator
DATE_LINE_RE = re.compile(r"^\s*(\d{1,2}/\d{1,2}/\d{2,4})\s*[/–-]\s*(.+?)\s*$")

# Shank depths lines; robust to "->" or HTML-like "-&gt;" and optional text between
SHANK_DEPTH_RE = re.compile(r"^shank\s*(\d+).*?(?:->|-&gt;)\s*([0-9?]+)", re.IGNORECASE)

# Numeric extractors
UM_NUM_RE = re.compile(r"(\d+)\s*(?:µm|um)?", re.IGNORECASE)
INT_RE = re.compile(r"\b(\d+)\b")

# Time extractor: accept colon or period, with optional AM/PM (case-insensitive)
# Examples matched: "1:16pm", "5.22", "13:20", "2:31 PM"
TIME_RE = re.compile(r"\b(\d{1,2}[:\.]\d{2})(?:\s*(am|pm))?\b", re.IGNORECASE)

# Specific regex for "Depth shank 4 entered relative to shank 1 ... 450[um]"
RELATIVE_SHANK4_RE = re.compile(
    r"Depth\s+shank\s*4\s+entered\s+relative\s+to\s+shank\s*1.*?(\d{2,5})\s*(?:um|µm)?",
    re.IGNORECASE,
)


# =============================================================================
# ----------------------- Helpers ----------------------------------------------
# =============================================================================
def canon(s: str) -> str:
    """
    Canonicalize a string for resilient header matching.

    Lowercases, strips, removes punctuation, and normalises unicode.

    Parameters
    ----------
    s : str
        Raw string to canonicalize.

    Returns
    -------
    str
        Canonical form with collapsed whitespace and no punctuation.
    """
    if not s:
        return ""
    s = unicodedata.normalize("NFKC", s).strip().lower()
    s = re.sub(r"[,:/()µ\-]+", " ", s)  # remove punctuation-like chars (keeps words)
    s = re.sub(r"\s+", " ", s)  # collapse whitespace
    return s


def cell_text_robust(cell) -> str:
    """
    Return the text content of a docx table cell.

    Falls back to joining non-empty paragraph texts when ``cell.text`` is empty.

    Parameters
    ----------
    cell : docx.table._Cell
        A python-docx table cell object.

    Returns
    -------
    str
        Stripped cell text.
    """
    txt = cell.text.strip()
    if txt:
        return txt
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(parts)


def parse_value_multiline(text: str) -> List[str]:
    """
    Split *text* into non-empty stripped lines.

    Parameters
    ----------
    text : str
        Multi-line string to split.

    Returns
    -------
    list[str]
        List of non-empty stripped lines.
    """
    return [ln.strip() for ln in text.splitlines() if ln.strip()]


def parse_date_lines_from_text(block_text: str) -> List[Dict[str, str]]:
    """
    Extract date/experiment pairs from a block of text.

    Parameters
    ----------
    block_text : str
        Multi-line text block potentially containing lines of the form
        ``"MM/DD/YYYY - Experiment Name"``.

    Returns
    -------
    list[dict]
        List of ``{"date": str, "experiment": str}`` dicts for every matched line.
    """
    items = []
    for ln in parse_value_multiline(block_text):
        m = DATE_LINE_RE.match(ln)
        if m:
            items.append({"date": m.group(1), "experiment": m.group(2).strip()})
    return items


def normalize_um(value: str) -> int | None:
    """
    Parse a micrometre depth string to an integer.

    Parameters
    ----------
    value : str
        String potentially containing a numeric value followed by ``um`` or ``µm``
        (e.g. ``"450 um"`` or ``"450"``).

    Returns
    -------
    int or None
        Parsed integer, or *value* unchanged when it is not a string.
    """
    if not isinstance(value, str):
        return value
    m = UM_NUM_RE.search(value)
    return int(m.group(1)) if m else None


def normalize_int(value: str) -> int | None:
    """
    Parse the first integer found in a string.

    Parameters
    ----------
    value : str
        String to search for an integer token.

    Returns
    -------
    int or None
        First integer found, or *value* unchanged when it is not a string.
    """
    if not isinstance(value, str):
        return value
    m = INT_RE.search(value)
    return int(m.group(1)) if m else None


def extract_time(text: str) -> str | None:
    """
    Extract time from text supporting colon or period separator and optional AM/PM.
    IMPORTANT: Returns the matched time string EXACTLY as it appears (no normalization).
    Examples returned: '1:16pm', '5.22', '2:31 PM'
    """
    m = TIME_RE.search(text)
    if not m:
        return None
    # Return the exact matched substring (includes AM/PM if present)
    return m.group(0)


def is_date_experiment_header(kcanon: str) -> bool:
    """
    Detect ``Date / Experiment`` table headers robustly.

    Handles extra text such as ``"(delete irrelevant fields)"``.

    Parameters
    ----------
    kcanon : str
        Canonicalized header string.

    Returns
    -------
    bool
        ``True`` if *kcanon* represents a Date/Experiment header.
    """
    if "date experiment" in kcanon:
        return True
    return "date" in kcanon and "experiment" in kcanon


def is_known_field_start(kcanon: str) -> bool:
    """
    Return whether *kcanon* maps to a known metadata field (excluding ``date_experiment``).

    Used to detect the end of a date/experiment multi-row block.

    Parameters
    ----------
    kcanon : str
        Canonicalized header string.

    Returns
    -------
    bool
        ``True`` if *kcanon* matches a known field other than ``date_experiment``.
    """
    return kcanon in KNOWN_FIELDS_CANON and KNOWN_FIELDS_CANON[kcanon] != "date_experiment"


def resolve_note_key_from_canon(kcanon: str) -> str | None:
    """
    Return the notes data-dict key for a given canonical header, or ``None``.

    Parameters
    ----------
    kcanon : str
        Canonicalized header string.

    Returns
    -------
    str or None
        Data key such as ``"surgery_notes"`` or ``"metabond_notes"``, or ``None``
        if *kcanon* does not represent a notes section.
    """
    if kcanon in NOTE_HEADERS_CANON:
        return NOTE_HEADERS_CANON[kcanon]
    for note_hdr_canon, data_key in NOTE_HEADERS_CANON.items():
        if kcanon.startswith(note_hdr_canon) or note_hdr_canon in kcanon:
            return data_key
    return None


def strip_note_header(line: str) -> str:
    """
    Remove a note header prefix from the given line and return trailing content only.
    If no header is present, returns the line unchanged.
    """
    for pat in NOTE_HEADER_PATTERNS.values():
        m = pat.match(line)
        if m:
            return m.group(1).strip()
    return line


def append_note_lines(data: Dict[str, Any], key: str, lines: List[str]) -> None:
    """
    Append lines to a notes field without duplicates, stripping header prefixes.

    Preserves order of first occurrence.

    Parameters
    ----------
    data : dict
        Mutable data dictionary to update in-place.
    key : str
        Notes field key (e.g. ``"surgery_notes"``).
    lines : list[str]
        Lines to append.

    Returns
    -------
    None
    """
    existing_lines = [ln.strip() for ln in data.get(key, "").splitlines() if ln.strip()]
    seen = set(existing_lines)
    for ln in lines:
        cleaned = strip_note_header(ln).strip()
        if not cleaned:
            continue
        if cleaned not in seen:
            existing_lines.append(cleaned)
            seen.add(cleaned)
    data[key] = "\n".join(existing_lines)


# =============================================================================
# ----------------------- Parsers ----------------------------------------------
# =============================================================================


def parse_metadata_table(
    doc: Document, table_index: int = 0
) -> Tuple[Dict[str, Any], pd.DataFrame]:
    """
    Parse the main surgery-notes metadata table.

    Handles a multi-row Date/Experiment block and notes embedded in table cells.

    Parameters
    ----------
    doc : Document
        Loaded python-docx Document object.
    table_index : int, optional
        Index of the metadata table within *doc*. Default is ``0``.

    Returns
    -------
    tuple[dict, pd.DataFrame]
        ``(data, date_expt_df)`` where *data* is the structured field dict and
        *date_expt_df* has columns ``["date", "experiment"]``.
    """
    table = doc.tables[table_index]

    data: Dict[str, Any] = {
        "date_experiment": [],
        "goal": "",
        "overall_impressions": "",
        "animal_id": None,
        "anesthesia_induction": "",
        "heating_pad": "",
        "animal_birthdate": "",
        "weight_before": "",
        "weight_after": "",
        "surgery_history": "",
        "craniotomy_coordinates": "",
        "targeting_modifications": [],
        "final_probe_depth_um": None,
        "z_gnd_screw_ref_screw": "",
        "z_gnd_probe_ref_probe": "",
        "probe_sn": "",
        "probe_gnd_ref_connected_to": "",
        "probe_orientation_ap": "",
        "probe_channels": None,
        # Paragraph-only fields will be filled later (initialized here for unified dict)
        "entered_brain_time": None,
        "depth_shank4_relative_um": None,
        "entered_dead_zone": {},
        "entered_piriform": {},
        "final_depth_um": None,
        "hit_final_depth_time": None,
        "metabond_applied_time": None,
        "relyx_started_time": None,
        "surgery_end_time": None,
        "surgery_notes": "",
        "metabond_notes": "",
        "probe_notes": "",
        "penetration_notes": "",
        "implantation_notes": "",
        "plugin_notes": "",
    }

    in_date_block = False

    for row in table.rows:
        cells = row.cells
        if len(cells) < 1:
            continue

        raw_key = cell_text_robust(cells[0])
        key_canon = canon(raw_key)
        raw_val = "\n".join(cell_text_robust(c) for c in cells[1:]).strip()

        # Date/Experiment block collection across subsequent rows (table-only)
        if in_date_block:
            if is_known_field_start(key_canon) or resolve_note_key_from_canon(key_canon):
                in_date_block = False
            else:
                combined = [cell_text_robust(cells[0])]
                if len(cells) > 1:
                    combined.extend(cell_text_robust(c) for c in cells[1:])
                items = parse_date_lines_from_text("\n".join(t for t in combined if t.strip()))
                for it in items:
                    if it not in data["date_experiment"]:
                        data["date_experiment"].append(it)
                continue

        # Skip pure placeholder unless it's part of Date/Experiment header
        if "delete irrelevant" in key_canon and not is_date_experiment_header(key_canon):
            note_key = resolve_note_key_from_canon(key_canon)
            if note_key:
                append_note_lines(data, note_key, parse_value_multiline(raw_val))
            continue

        # Start Date/Experiment header
        if is_date_experiment_header(key_canon):
            if raw_val:
                data["date_experiment"].extend(parse_date_lines_from_text(raw_val))
            in_date_block = True
            continue

        # Notes in table rows
        note_key = resolve_note_key_from_canon(key_canon)
        if note_key:
            append_note_lines(data, note_key, parse_value_multiline(raw_val))
            continue

        # Normal known-field parsing
        if raw_key in KNOWN_FIELDS:
            field = KNOWN_FIELDS[raw_key]

            if field == "targeting_modifications":
                data[field] = parse_value_multiline(raw_val)
            elif field == "final_probe_depth_um":
                data[field] = normalize_um(raw_val)
            elif field == "probe_channels":
                data[field] = normalize_int(raw_val)
            elif field == "animal_id":
                id_match = re.search(r"\b(\d{4,})\b", raw_val) or re.search(
                    r"\b(\d{4,})\b", raw_key
                )
                data[field] = id_match.group(1) if id_match else (raw_val or raw_key)
            else:
                data[field] = "\n".join(parse_value_multiline(raw_val))
        else:
            # Canonical fallback (excluding date_experiment)
            if (
                key_canon in KNOWN_FIELDS_CANON
                and KNOWN_FIELDS_CANON[key_canon] != "date_experiment"
            ):
                field = KNOWN_FIELDS_CANON[key_canon]
                if field == "targeting_modifications":
                    data[field] = parse_value_multiline(raw_val)
                elif field == "final_probe_depth_um":
                    data[field] = normalize_um(raw_val)
                elif field == "probe_channels":
                    data[field] = normalize_int(raw_val)
                elif field == "animal_id":
                    id_match = re.search(r"\b(\d{4,})\b", raw_val) or re.search(
                        r"\b(\d{4,})\b", raw_key
                    )
                    data[field] = id_match.group(1) if id_match else (raw_val or raw_key)
                else:
                    data[field] = "\n".join(parse_value_multiline(raw_val))
            # else: ignore unknown keys for schema cleanliness

    date_expt_df = (
        pd.DataFrame(data["date_experiment"])
        if data["date_experiment"]
        else pd.DataFrame(columns=["date", "experiment"])
    )
    return data, date_expt_df


def parse_paragraph_sections(doc: Document, data: Dict[str, Any]) -> None:
    """
    Parse times, depths, and free-text notes from document paragraphs.

    Iterates per physical line to prevent duplicates and capture inline header
    content cleanly. Does NOT touch table content. Times are returned exactly
    as matched (no normalisation).

    Parameters
    ----------
    doc : Document
        Loaded python-docx Document object.
    data : dict
        Mutable data dictionary (from :func:`parse_metadata_table`) updated in-place.

    Returns
    -------
    None
    """
    current_notes_key = None
    in_dead_zone = False
    in_piriform = False

    for para in doc.paragraphs:
        # Iterate per physical line inside the paragraph to separate mixed content
        for raw_line in para.text.splitlines():
            line = raw_line.strip()
            if not line:
                continue

            # Normalize common artifacts before matching (HTML-like "->")
            line_norm = line.replace("-&gt;", "->").replace("→", "->")
            lc = canon(line_norm)  # for header detection only

            # End notes block when K/X anesthesia section starts
            if lc.startswith("k x anesthesia"):
                current_notes_key = None
                in_dead_zone = False
                in_piriform = False
                continue

            # Notes header with inline content? Capture header + trailing content immediately
            appended_inline = False
            for note_key, pat in NOTE_HEADER_PATTERNS.items():
                m = pat.match(line)  # use raw line to preserve punctuation and spacing
                if m:
                    current_notes_key = note_key
                    trailing = m.group(1).strip()
                    if trailing:
                        append_note_lines(data, current_notes_key, [trailing])
                        appended_inline = True
                    break
            if not appended_inline:
                # Fallback notes header detection (canonical startswith/contains)
                nk = resolve_note_key_from_canon(lc)
                if nk and nk != current_notes_key:
                    current_notes_key = nk
                    in_dead_zone = False
                    in_piriform = False
                    # header line without inline content; proceed

            # Dead zone / Piriform headers (end notes)
            if "(for apc) entered dead zone" in lc:
                in_dead_zone = True
                in_piriform = False
                current_notes_key = None
                continue

            if "(for apc) entered piriform" in lc:
                in_piriform = True
                in_dead_zone = False
                current_notes_key = None
                continue

            # Collect shank depths when in a zone section (paragraph-only)
            if in_dead_zone or in_piriform:
                # Accept 'Never' and values ending with '?'
                m = SHANK_DEPTH_RE.match(line_norm)
                if m:
                    shank = int(m.group(1))
                    val = m.group(2)
                    if isinstance(val, str) and val.strip().lower() == "never":
                        depth = None
                    else:
                        # strip any non-digits like '?'
                        digits = re.search(r"(\d+)", val)
                        depth = int(digits.group(1)) if digits else None
                    bucket = "entered_dead_zone" if in_dead_zone else "entered_piriform"
                    data[bucket][shank] = depth
                continue

            # Times & scalar fields (paragraph-only) — match PHRASES on raw line, extract time with TIME_RE
            if "Entered brain at" in line_norm:
                t = extract_time(line_norm)
                if t:
                    data["entered_brain_time"] = t

            m_rel = RELATIVE_SHANK4_RE.search(line_norm)
            if m_rel:
                data["depth_shank4_relative_um"] = int(m_rel.group(1))

            if line_norm.upper().startswith("FINAL DEPTH"):
                m = re.search(r"(\d+)", line_norm)
                if m:
                    data["final_depth_um"] = int(m.group(1))

            if "Hit final depth at" in line_norm:
                t = extract_time(line_norm)
                if t:
                    data["hit_final_depth_time"] = t

            if "Metabond applied at" in line_norm or "Bruno cement applied at" in line_norm:
                t = extract_time(line_norm)
                if t:
                    data["metabond_applied_time"] = t

            if "RelyX started at" in line_norm:
                t = extract_time(line_norm)
                if t:
                    data["relyx_started_time"] = t

            if "Surgery ended at" in line_norm:
                t = extract_time(line_norm)
                if t:
                    data["surgery_end_time"] = t

            # Append the line to the active note (if any), once and header-stripped.
            if current_notes_key:
                if not appended_inline:
                    append_note_lines(data, current_notes_key, [line])
                continue
            # else: ignore non-note lines that don't belong to recognized sections


def parse_kx_anesthesia_table(doc: Document) -> pd.DataFrame:
    """
    Parse the K/X anesthesia dosing table into a DataFrame.

    Robust to rows with missing columns and note-only lines.

    Parameters
    ----------
    doc : Document
        Loaded python-docx Document object.

    Returns
    -------
    pd.DataFrame
        DataFrame with columns ``["time", "agent", "concentration", "volume_ul", "reason"]``.
        Empty DataFrame with those columns is returned when no matching table is found.
    """
    kx_rows: List[Dict[str, Any]] = []

    def is_kx_header_row(headers: List[str]) -> bool:
        hs = [h.strip().lower() for h in headers]
        return "time" in hs and "agent" in hs and "concentration" in hs

    # Find a table whose first row looks like the K/X header
    for table in doc.tables:
        if len(table.rows) == 0:
            continue

        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if not header_cells:
            continue

        if is_kx_header_row(header_cells):
            # Parse table rows
            last_row_dict = None
            for r_idx, row in enumerate(table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                row_dict: Dict[str, Any] = {}
                if len(cells) >= 1:
                    row_dict["time"] = cells[0] or None
                if len(cells) >= 2:
                    row_dict["agent"] = cells[1] or None
                if len(cells) >= 3:
                    row_dict["concentration"] = cells[2] or None
                if len(cells) >= 4:
                    row_dict["volume_ul"] = cells[3] or None
                if len(cells) >= 5:
                    row_dict["reason"] = cells[4] or None

                # Note-only lines (single cell after a data row)
                if len(cells) == 1 and cells[0] and last_row_dict:
                    if not last_row_dict.get("reason"):
                        last_row_dict["reason"] = cells[0]
                    else:
                        kx_rows.append(
                            {
                                "time": None,
                                "agent": None,
                                "concentration": None,
                                "volume_ul": None,
                                "reason": cells[0],
                            }
                        )
                    continue

                # Skip truly empty rows
                if not any(cells):
                    continue

                kx_rows.append(row_dict)
                last_row_dict = row_dict
            break  # stop after the first matching table

    kx_df = (
        pd.DataFrame(kx_rows)
        if kx_rows
        else pd.DataFrame(columns=["time", "agent", "concentration", "volume_ul", "reason"])
    )
    return kx_df


# =============================================================================
# ----------------------- DOCX parsing -----------------------------------------
# =============================================================================
def parse_surgery_notes(
    surgery_notes_path: Path,
    metadata_table_index: int = 0,
) -> Tuple[Dict[str, Any], pd.DataFrame, pd.DataFrame]:
    """
    End-to-end parser for a surgery notes DOCX file.

    Parses the metadata table (including Date/Experiment block and notes-in-table),
    paragraph sections (times, depths, free-text notes), and the K/X anesthesia table.
    Times are returned exactly as matched — no normalisation applied.

    Parameters
    ----------
    surgery_notes_path : Path
        Path to the ``.docx`` surgery notes file.
    metadata_table_index : int, optional
        Index of the main metadata table within the document. Default is ``0``.

    Returns
    -------
    tuple[dict, pd.DataFrame, pd.DataFrame, str | None]
        ``(data, date_expt_df, kx_df, probe_serial_number)`` where *data* contains
        all structured fields, *date_expt_df* has columns ``["date", "experiment"]``,
        *kx_df* has columns ``["time", "agent", "concentration", "volume_ul", "reason"]``,
        and *probe_serial_number* is the ``probe_sn`` field value (may be ``None``).

    Raises
    ------
    FileNotFoundError
        If *surgery_notes_path* does not exist.
    ValueError
        If the document contains no tables.
    """
    if not surgery_notes_path.exists():
        raise FileNotFoundError(f"Surgery notes file not found: {surgery_notes_path}")
    doc = Document(str(surgery_notes_path))
    if len(doc.tables) == 0:
        raise ValueError("No tables found in the document.")

    data, date_expt_df = parse_metadata_table(doc, table_index=metadata_table_index)
    parse_paragraph_sections(doc, data)
    kx_df = parse_kx_anesthesia_table(doc)

    # (Optional) unify final depth fields when table depth missing
    if data.get("final_probe_depth_um") is None and data.get("final_depth_um") is not None:
        data["final_probe_depth_um"] = data["final_depth_um"]
    probe_serial_number = data.get("probe_sn")

    return data, date_expt_df, kx_df, probe_serial_number


# =============================================================================
# ----------------------- Procedures builder ----------------------------------
# =============================================================================
def create_procedures_metadata(
    *,
    current_experiment: str,
    subject_id: str,
    protocol_id: str,
    surgeons: List[str],
    surgery_notes_path: Path,
    probe_device,
    probe_config,
) -> tuple[Procedures, str]:
    """
    Create AIND Procedures metadata from parsed surgery notes.

    Parameters
    ----------
    current_experiment : str
        One of ``{"delphi", "pirouette", "delphi_pirouette"}``. Only ``"pirouette"``
        triggers full surgery-note parsing; other values produce a minimal Procedures object.
    subject_id : str
        Unique subject identifier.
    protocol_id : str
        Protocol identifier applied to all Surgery records.
    surgeons : list[str]
        Names of surgeons. Index 0 is the primary surgeon; index 1 is used for
        the headframe surgery.
    surgery_notes_path : Path
        Path to the ``.docx`` surgery notes file.
    probe_device : EphysProbe
        Probe device object used in the ProbeImplant procedure.
    probe_config : ProbeConfig
        Probe configuration object used in the ProbeImplant procedure.

    Returns
    -------
    Procedures
        Validated AIND Procedures object containing craniotomy, implantation, and
        headframe Surgery records (or a minimal stub for non-pirouette experiments).
    """
    if "pirouette" in current_experiment:
        parsed_data, date_expt_df, kx_df, _ = parse_surgery_notes(surgery_notes_path)

        """Get craniotomy information"""
        # craniotomy starting weight
        match = re.search(r"Craniotomy:\s*([\d\.]+)", parsed_data.get("weight_before"))
        craniotomy_weight_before = float(match.group(1)) if match else None

        # craniotomy end weight
        match = re.search(r"Craniotomy:\s*([\d\.]+)", parsed_data.get("weight_after"))
        craniotomy_weight_after = float(match.group(1)) if match else None

        # induction type
        match = re.search(r"Craniotomy:\s*([^\s/]+)", parsed_data.get("anesthesia_induction"))
        craniotomy_induction = match.group(1) if match else None
        craniotomy_anesthetic = Anaesthetic(
            anaesthetic_type=craniotomy_induction,
            duration=210.0,  # hardcoded because a cranotiomy end time doesn't exist in the surgery notes
            duration_unit=TimeUnit.M,
        )

        craniotomy = Craniotomy(
            craniotomy_type=CraniotomyType.OTHER,
            protective_material=ProtectiveMaterial.KWIK_CAST,
            dura_removed=True,
        )

        # Create craniotomy object
        craniotomy_surgery = Surgery(
            protocol_id=protocol_id,
            start_date=pd.to_datetime(
                date_expt_df.loc[date_expt_df["experiment"] == "Craniotomy", "date"]
            ).iloc[0],
            experimenters=surgeons,
            animal_weight_prior=craniotomy_weight_before,
            animal_weight_post=craniotomy_weight_after,
            weight_unit=MassUnit.G,
            anaesthesia=craniotomy_anesthetic,
            procedures=[craniotomy],
            notes=parsed_data.get("surgery_notes"),
        )

        """Get implantation information"""
        # implantation starting weight
        match = re.search(r"Implantation:\s*([\d\.]+)", parsed_data.get("weight_before"))
        implantation_weight_before = float(match.group(1)) if match else None

        # implantation end weight
        match = re.search(r"Implantation:\s*([\d\.]+)", parsed_data.get("weight_after"))
        implantation_weight_after = float(match.group(1)) if match else None

        # 1) Keep only valid HH:MM entries
        valid = [t for t in kx_df["time"] if isinstance(t, str) and re.match(r"^\d{1,2}:\d{2}$", t)]

        # 2) Convert to (hour, minute) ints
        hm = [(int(h), int(m)) for h, m in (s.split(":") for s in valid)]

        # 3) Identify values strictly greater than 12:00.
        #    Rule: hours == 12 with minute > 0 are after noon; any entries AFTER the first 12:xx
        #    with hour in 1..11 are PM (due to sequence order).
        minutes_in_day = []  # convert to minutes since midnight
        seen_noon = False
        for s, (h, m) in zip(valid, hm):
            if h == 12 and m > 0:
                # Noon boundary
                minutes_in_day.append(h * 60 + m)
                seen_noon = True
            elif seen_noon and 1 <= h <= 11:
                # After noon
                minutes_in_day.append((h + 12) * 60 + m)
            else:
                # Before noon
                minutes_in_day.append(h * 60 + m)

        implantation_induction_duration = np.max(minutes_in_day) - np.min(minutes_in_day)

        # implantation anesthetic
        implantation_anesthetic = Anaesthetic(
            anaesthetic_type="".join(kx_df["agent"].unique()),
            duration=implantation_induction_duration,  # hardcoded because a cranotiomy end time doesn't exist in the surgery notes
            duration_unit=TimeUnit.M,
        )

        probe_implant = ProbeImplant(
            implanted_device=probe_device,
            device_config=probe_config,
        )

        # Create implantation object
        implantation_surgery = Surgery(
            protocol_id=protocol_id,
            start_date=pd.to_datetime(
                date_expt_df.loc[date_expt_df["experiment"] == "Implantation", "date"]
            ).iloc[0],
            experimenters=surgeons,
            animal_weight_prior=implantation_weight_before,
            animal_weight_post=implantation_weight_after,
            weight_unit=MassUnit.G,
            anaesthesia=implantation_anesthetic,
            procedures=[probe_implant],
            notes=parsed_data.get("implantation_notes"),
        )

        # Create headframe object
        headframe = Headframe(
            headframe_type="3D printed custom headframe for chronic implantation",
            headframe_part_number="CAD model: Headplate_mark12",
        )

        # Headframe surgery date
        match = re.search(r"(\d{1,2}/\d{1,2}/\d{2,4})", parsed_data.get("surgery_history"))
        date_str = match.group(1)
        for _fmt in ("%m/%d/%Y", "%m/%d/%y"):
            try:
                headframe_surgery_date = datetime.strptime(date_str, _fmt)
                break
            except ValueError:
                continue
        else:
            raise ValueError(f"Could not parse headframe surgery date: {date_str!r}")

        headframe_surgery = Surgery(
            protocol_id=protocol_id,
            start_date=headframe_surgery_date,
            experimenters=[surgeons[1]],
            procedures=[headframe],
        )

        # Create procedures object from surgery notes
        procedures = Procedures(
            subject_id=subject_id,
            subject_procedures=[craniotomy_surgery, implantation_surgery, headframe_surgery],
            coordinate_system=CoordinateSystemLibrary.BREGMA_ARI,
            notes=parsed_data.get("overall_impressions"),
        )
    else:
        procedures = Procedures(
            subject_id=subject_id,
        )

    return procedures
